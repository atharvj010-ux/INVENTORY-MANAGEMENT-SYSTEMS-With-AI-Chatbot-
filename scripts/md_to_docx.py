"""Convert PROJECT_DOCUMENTATION.md to Word (.docx)."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(cols):
            text = row[ci].strip() if ci < len(row) else ""
            text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
            cell = table.rows[ri].cells[ci]
            cell.text = text
            if ri == 0:
                set_cell_shading(cell, "1F2937")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
    doc.add_paragraph()


def add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(40, 40, 40)


def convert(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Nexus AI — Inventory Management App", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    sub = doc.add_paragraph("Complete Project Documentation")
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub.runs[0].italic = True
    doc.add_paragraph()

    i = 0
    in_code = False
    code_buf: list[str] = []
    table_buf: list[list[str]] = []

    def flush_table() -> None:
        nonlocal table_buf
        if table_buf:
            add_table(doc, table_buf)
            table_buf = []

    while i < len(lines):
        line = lines[i]

        if in_code:
            if line.strip().startswith("```"):
                in_code = False
                add_code_block(doc, code_buf)
                code_buf = []
            else:
                code_buf.append(line)
            i += 1
            continue

        if line.strip().startswith("```"):
            flush_table()
            in_code = True
            code_buf = []
            i += 1
            continue

        if "|" in line and line.strip().startswith("|"):
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            if cells and not all(re.match(r"^[-:]+$", c) for c in cells):
                table_buf.append(cells)
            i += 1
            continue
        else:
            flush_table()

        if line.strip() == "---":
            doc.add_paragraph()
            i += 1
            continue

        m = re.match(r"^(#{1,4})\s+(.+)$", line)
        if m:
            level = len(m.group(1))
            heading_text = re.sub(r"\*\*([^*]+)\*\*", r"\1", m.group(2))
            doc.add_heading(heading_text, level=min(level, 4))
            i += 1
            continue

        if line.strip().startswith("- ") or line.strip().startswith("* "):
            item = line.strip()[2:]
            item = re.sub(r"\*\*([^*]+)\*\*", r"\1", item)
            item = re.sub(r"`([^`]+)`", r"\1", item)
            doc.add_paragraph(item, style="List Bullet")
            i += 1
            continue

        if re.match(r"^\d+\.\s", line.strip()):
            item = re.sub(r"^\d+\.\s", "", line.strip())
            item = re.sub(r"\*\*([^*]+)\*\*", r"\1", item)
            doc.add_paragraph(item, style="List Number")
            i += 1
            continue

        if line.strip():
            para_text = line.strip()
            para_text = re.sub(r"\*\*([^*]+)\*\*", r"\1", para_text)
            para_text = re.sub(r"`([^`]+)`", r"\1", para_text)
            para_text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", para_text)
            p = doc.add_paragraph()
            parts = re.split(r"(\*\*[^*]+\*\*)", para_text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        i += 1

    flush_table()
    doc.save(docx_path)
    print(f"Created: {docx_path}")


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    md = root / "PROJECT_DOCUMENTATION.md"
    out = root / "PROJECT_DOCUMENTATION.docx"
    if len(sys.argv) > 1:
        out = Path(sys.argv[1])
    convert(md, out)
